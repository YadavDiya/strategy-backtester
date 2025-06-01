from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def create_formatted_docx():
    # Create a new Document
    doc = Document()
    
    # Set the margins
    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    # Title
    title = doc.add_heading('Strategy Backtester - User Manual', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Table of Contents
    doc.add_heading('Table of Contents', 1)
    toc_items = [
        'Introduction',
        'Installation',
        'Quick Start',
        'Features',
        'Usage Guide',
        'Output Files',
        'Troubleshooting',
        'Support'
    ]
    for idx, item in enumerate(toc_items, 1):
        p = doc.add_paragraph()
        p.add_run(f"{idx}. {item}")
    
    doc.add_page_break()

    # Introduction
    doc.add_heading('Introduction', 1)
    doc.add_paragraph(
        'The Strategy Backtester is a powerful Python-based tool for testing '
        'cryptocurrency trading strategies using historical data from Binance. '
        'It allows you to evaluate trading strategies before risking real capital.'
    )

    # Key Features
    doc.add_heading('Key Features', 2)
    features = [
        'Real-time data fetching from Binance',
        'Multiple pre-built trading strategies',
        'Detailed performance metrics',
        'Excel-formatted reports with trade analysis',
        'Easy strategy customization'
    ]
    for feature in features:
        p = doc.add_paragraph()
        p.style = 'List Bullet'
        p.add_run(feature)

    # Installation
    doc.add_heading('Installation', 1)
    
    # Prerequisites
    doc.add_heading('Prerequisites', 2)
    prereqs = ['Python 3.x', 'Git']
    for prereq in prereqs:
        p = doc.add_paragraph()
        p.style = 'List Bullet'
        p.add_run(prereq)

    # Setup Steps
    doc.add_heading('Setup Steps', 2)
    
    # Code blocks with different formatting
    code_blocks = [
        ('Clone the repository:', 'git clone [repository-url]\ncd strategy-backtester'),
        ('Create a virtual environment:', '# Windows\npython -m venv venv\n.\\venv\\Scripts\\activate\n\n# Linux/Mac\npython3 -m venv venv\nsource venv/bin/activate'),
        ('Install dependencies:', 'pip install -r requirements.txt')
    ]
    
    for idx, (title, code) in enumerate(code_blocks, 1):
        p = doc.add_paragraph()
        p.add_run(f"{idx}. {title}")
        p = doc.add_paragraph()
        p.style = 'No Spacing'
        code_run = p.add_run(code)
        code_run.font.name = 'Courier New'
        code_run.font.size = Pt(10)
        
    # Quick Start
    doc.add_heading('Quick Start', 1)
    steps = [
        'Activate your virtual environment (if not already activated)',
        'Run the backtester:\n```bash\npython main.py\n```',
        'Find results in the `backtester_output_[DATE]` directory'
    ]
    for idx, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.add_run(f"{idx}. {step}")

    # Features
    doc.add_heading('Features', 1)
    
    # Implemented Strategies
    doc.add_heading('Implemented Strategies', 2)
    
    strategies = {
        'MACD Strategy': [
            'Uses MACD line crossing its EMA',
            'Configurable parameters for different timeframes',
            'Suitable for trend following'
        ],
        'RSI-EMA Strategy': [
            'Combines RSI and EMA indicators',
            'Entry: RSI > 30 and price > EMA(21)',
            'Exit: RSI > 70 or price < EMA(21)'
        ]
    }
    
    for strategy, points in strategies.items():
        doc.add_heading(strategy, 3)
        for point in points:
            p = doc.add_paragraph()
            p.style = 'List Bullet'
            p.add_run(point)

    # Technical Indicators
    doc.add_heading('Technical Indicators', 2)
    indicators = [
        'EMA (Exponential Moving Average)',
        'MACD (Moving Average Convergence Divergence)',
        'RSI (Relative Strength Index)'
    ]
    for indicator in indicators:
        p = doc.add_paragraph()
        p.style = 'List Bullet'
        p.add_run(indicator)

    # Usage Guide
    doc.add_heading('Usage Guide', 1)
    
    # Basic Usage
    doc.add_heading('Basic Usage', 2)
    p = doc.add_paragraph('Run the backtester with default settings:')
    p = doc.add_paragraph()
    code_run = p.add_run('python main.py')
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(10)

    # Output Files
    doc.add_heading('Output Files', 1)
    
    # Location
    doc.add_heading('Location', 2)
    p = doc.add_paragraph('All output files are saved in a directory named:')
    p = doc.add_paragraph()
    code_run = p.add_run('backtester_output_YYYY_MM_DD/')
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(10)

    # File Types
    doc.add_heading('File Types', 2)
    file_types = {
        'Trade Results (backtester_results_YYYYMMDD_HHMMSS.xlsx)': [
            'Entry Time',
            'Entry Price',
            'Exit Time',
            'Exit Price',
            'Strategy Name',
            'PnL (Profit/Loss)',
            'Trade Status'
        ],
        'Summary Metrics (summary_metrics_YYYYMMDD_HHMMSS.xlsx)': [
            'Total Trades',
            'Win Rate',
            'Total PnL',
            'Average PnL'
        ]
    }
    
    for file_type, fields in file_types.items():
        doc.add_heading(file_type, 3)
        for field in fields:
            p = doc.add_paragraph()
            p.style = 'List Bullet'
            p.add_run(field)

    # Troubleshooting
    doc.add_heading('Troubleshooting', 1)
    
    # Common Issues
    doc.add_heading('Common Issues', 2)
    issues = {
        'Installation Problems': [
            'Ensure Python 3.x is installed',
            'Use a fresh virtual environment',
            'Update pip: pip install --upgrade pip'
        ],
        'Data Fetching Issues': [
            'Check internet connection',
            'Verify Binance API availability',
            'Ensure symbol names are correct'
        ],
        'Excel Output Errors': [
            'Install openpyxl: pip install openpyxl',
            'Close any open Excel files',
            'Check write permissions in output directory'
        ]
    }
    
    for issue, solutions in issues.items():
        doc.add_heading(issue, 3)
        for solution in solutions:
            p = doc.add_paragraph()
            p.style = 'List Bullet'
            p.add_run(solution)

    # Support
    doc.add_heading('Support', 1)
    support_steps = [
        'Check the [Issues] section on GitHub',
        'Create a new issue with detailed description',
        'Include system information and error messages'
    ]
    for step in support_steps:
        p = doc.add_paragraph()
        p.style = 'List Bullet'
        p.add_run(step)

    # Footnote
    doc.add_page_break()
    doc.add_heading('Footnote: Suggested Improvements', 1)
    
    # Immediate Enhancements
    doc.add_heading('Immediate Enhancements', 2)
    enhancements = {
        'Risk Management Features': [
            'Implement stop-loss and take-profit functionality',
            'Add position sizing based on risk percentage',
            'Include maximum drawdown calculations'
        ],
        'Advanced Analytics': [
            'Add visualization of trade entries/exits',
            'Include equity curve analysis',
            'Calculate Sharpe ratio and other risk metrics'
        ],
        'Strategy Enhancements': [
            'Add support for multiple timeframes',
            'Implement portfolio backtesting',
            'Add support for short positions'
        ],
        'User Interface': [
            'Create a web-based dashboard',
            'Add real-time strategy monitoring',
            'Include interactive parameter optimization'
        ],
        'Data Management': [
            'Add support for multiple data sources',
            'Implement data caching',
            'Add custom timeframe resampling'
        ]
    }
    
    for category, items in enhancements.items():
        doc.add_heading(category, 3)
        for item in items:
            p = doc.add_paragraph()
            p.style = 'List Bullet'
            p.add_run(item)

    # Future Roadmap
    doc.add_heading('Future Roadmap', 2)
    versions = {
        'Version 2.0': [
            'Machine learning strategy integration',
            'Real-time paper trading mode',
            'Strategy optimization using genetic algorithms'
        ],
        'Version 3.0': [
            'Live trading integration',
            'Mobile app for monitoring',
            'Cloud-based backtesting'
        ]
    }
    
    for version, features in versions.items():
        doc.add_heading(version, 3)
        for feature in features:
            p = doc.add_paragraph()
            p.style = 'List Bullet'
            p.add_run(feature)

    # Save the document
    doc.save('docs/User Manual.docx')

if __name__ == '__main__':
    create_formatted_docx() 